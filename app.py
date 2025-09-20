import streamlit as st
from openai import OpenAI
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from docx import Document

# --- OpenAI API Key (from Streamlit Secrets) ---
client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

# --- Streamlit Page Config ---
st.set_page_config(page_title="ETL to PySpark Validator", page_icon="⚡", layout="wide")

# ----------- Custom Styling -----------
st.markdown(
    """
    <style>
    .main {
        background-color: #f9fafb;
    }
    .stButton>button {
        width: 100%;
        border-radius: 8px;
        height: 3em;
        font-size: 16px;
        font-weight: bold;
        background: linear-gradient(90deg, #2563eb, #1d4ed8);
        color: white;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

st.title("⚡ ETL → PySpark Validator")
st.write("Upload **Informatica or Datastage file** along with the **PySpark output file** and validate the conversion using AI. The app will also generate a corrected PySpark file if needed.")

col1, col2 = st.columns(2)

# ----------- File Upload Section -----------
with col1:
    st.subheader("ETL Input File")
    etl_option = st.radio(
        "Choose ETL Type:",
        ["Informatica", "Datastage"],
        horizontal=True
    )

    if etl_option == "Informatica":
        informatica_file = st.file_uploader("Upload Informatica File", type=["xml", "json", "txt"], key="inf")
        datastage_file = None
    else:
        datastage_file = st.file_uploader("Upload Datastage File", type=["xml", "dsx", "txt"], key="ds")
        informatica_file = None

with col2:
    st.subheader("PySpark Output File")
    pyspark_file = st.file_uploader("Upload PySpark File", type=["py"], key="pyspark")

validation_report = None
corrected_pyspark = None

# ----------- Validation Section -----------
if st.button("🚀 Validate Conversion"):
    if (informatica_file or datastage_file) and pyspark_file:
        st.info("⏳ Validating conversion... please wait.")

        # Read contents
        input_content = ""
        if informatica_file:
            input_content = informatica_file.read().decode("utf-8", errors="ignore")
        elif datastage_file:
            input_content = datastage_file.read().decode("utf-8", errors="ignore")

        pyspark_content = pyspark_file.read().decode("utf-8", errors="ignore")

        # Call OpenAI API for validation
        validation_prompt = f"""
        You are an ETL to PySpark conversion validator.
        
        Input ETL file (from {'Informatica' if informatica_file else 'Datastage'}):
        {input_content}

        Output PySpark file:
        {pyspark_content}

        Validate whether the PySpark code correctly implements the ETL logic.
        Provide a detailed validation report with clear sections:
        - ✅ Correct parts
        - ⚠️ Potential issues
        - ❌ Missing logic
        - 💡 Suggested improvements
        """

        try:
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": validation_prompt}],
                temperature=0
            )

            validation_report = response.choices[0].message.content.strip()
            st.success("✅ Validation Completed")
            st.markdown("### 📝 Validation Report")
            st.write(validation_report)

            # --- Step 2: Ask LLM to correct the PySpark file ---
            correction_prompt = f"""
            Based on the ETL input and PySpark output above, rewrite the PySpark code so that it
            fully and correctly implements the ETL logic.
            
            IMPORTANT:
            - Return only the corrected PySpark code.
            - If the original file is already correct, return the same code unchanged.
            """

            correction_response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are an expert PySpark converter."},
                    {"role": "user", "content": f"ETL Input:\n{input_content}\n\nPySpark Output:\n{pyspark_content}\n\n{correction_prompt}"}
                ],
                temperature=0
            )

            corrected_pyspark = correction_response.choices[0].message.content.strip()

        except Exception as e:
            st.error(f"Error during validation: {e}")
    else:
        st.warning("⚠️ Please upload both an ETL file (Informatica/Datastage) and a PySpark file.")

# ----------- Helper Functions for Reports -----------
def parse_sections(text):
    sections = {"Correct Parts": [], "Potential Issues": [], "Missing Logic": [], "Suggested Improvements": []}
    current = None
    for line in text.split("\n"):
        line = line.strip()
        if "Correct parts" in line or "Correct Parts" in line:
            current = "Correct Parts"
        elif "Potential issues" in line or "Potential Issues" in line:
            current = "Potential Issues"
        elif "Missing logic" in line or "Missing Logic" in line:
            current = "Missing Logic"
        elif "Suggested improvements" in line or "Suggested Improvements" in line:
            current = "Suggested Improvements"
        elif line.startswith("-") or line.startswith("•"):
            if current:
                sections[current].append(line.lstrip("-• ").strip())
    return sections

def create_pdf(sections):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    elements = []

    elements.append(Paragraph("ETL → PySpark Validation Report", styles["Title"]))
    elements.append(Spacer(1, 12))

    for sec, items in sections.items():
        elements.append(Paragraph(sec, styles["Heading2"]))
        if items:
            bullet_list = ListFlowable(
                [ListItem(Paragraph(item, styles["Normal"])) for item in items],
                bulletType="bullet",
            )
            elements.append(bullet_list)
        else:
            elements.append(Paragraph("No findings.", styles["Normal"]))
        elements.append(Spacer(1, 12))

    doc.build(elements)
    buffer.seek(0)
    return buffer

def create_docx(sections):
    buffer = BytesIO()
    doc = Document()
    doc.add_heading("ETL → PySpark Validation Report", 0)

    for sec, items in sections.items():
        doc.add_heading(sec, level=1)
        if items:
            for item in items:
                doc.add_paragraph(item, style="List Bullet")
        else:
            doc.add_paragraph("No findings.")
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ----------- Download Section -----------
if validation_report:
    col1, col2 = st.columns(2)
    sections = parse_sections(validation_report)

    pdf_file = create_pdf(sections)
    docx_file = create_docx(sections)

    with col1:
        st.download_button(
            label="⬇️ Download Report (PDF)",
            data=pdf_file,
            file_name="Validation_Report.pdf",
            mime="application/pdf"
        )
    with col2:
        st.download_button(
            label="⬇️ Download Report (Word)",
            data=docx_file,
            file_name="Validation_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# ----------- Corrected PySpark Download -----------
if corrected_pyspark:
    st.markdown("### 🛠️ Corrected PySpark Code")
    st.code(corrected_pyspark, language="python")

    st.download_button(
        label="⬇️ Download Corrected PySpark File",
        data=corrected_pyspark.encode("utf-8"),
        file_name="Corrected_PySpark.py",
        mime="text/x-python"
    )
